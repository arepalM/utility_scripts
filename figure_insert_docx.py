# -*- coding: utf-8 -*-
"""
Created on Mon Apr 25 21:44:43 2022
@author: Mike Lapera

NOTES:

This script formats and places png files with captions all together in a single word document (one for each subfolder) in the parent directory

1. Be sure to place this python script in the parent directory above the subfolders for each model type (i.e. each of the subfolder in directory list, Line 31)
2. Each subfolder should have the png files you're looking to insert into the word doc
3. The variable "filename" must follow this convention: modelname_plottitle_RRTvalue_RRSvalue
4. If you change the file naming scheme, be sure to update figure_caption (Line 28) and NN_name, RTT, and RSS variables which are parsed from it

    

"""

import os
import docx 
from docx.shared import Inches
import glob
#from exceptions import PendingDeprecationWarning

# Figure number index
pic_index = 1
# Formatting for figure caption
figure_caption = 'Figure {}: {}, RTT = {}, RRS = {}'


directory = ['DNN1','DNN2','DNN3']
base = os.getcwd()
for i in range(len(directory)):
    document = docx.Document()
    folder = os.path.join(base,directory[i])
    os.chdir(path = folder)
    filelist = glob.glob('*.png')
    for j in range(len(filelist)):
        filename = filelist[j].split('_')
        NN_name = filename[0]
        RTT = '{}%'.format(filename[2][0:2])
        RSS = filename[3][0:3]
        #ip2.add_picture(filelist[j])
        document.add_picture(filelist[j],width=Inches(6),height=Inches(6))
        document.add_paragraph(figure_caption.format(pic_index,NN_name,RTT,RSS))
        os.chdir('..')
        document.save('{}_figures.docx'.format(NN_name))
        pic_index += 1
        os.chdir(folder)
    